import logging
import os
from typing import Dict, Any
from docx import Document
from pdfminer.high_level import extract_text
from sqlalchemy.orm import Session
from utils.structure import DataStructureService

class ResumeParserService:
    def __init__(self, db: Session):
        self.db = db
        self.data_structure_service = DataStructureService()

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse resume file and extract structured information"""
        file_ext = os.path.splitext(file_path)[1].lower()
        content = self._extract_text(file_path, file_ext)
        return self.data_structure_service.structure_resume_data(content)

    def _extract_text_from_file_obj(self, file_obj: Any, file_ext: str) -> str:
        """Extract text content from file object"""
        if file_ext == '.pdf':
            return extract_text(file_obj)
        elif file_ext in ['.docx', '.doc']:
            doc = Document(file_obj)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError(f'Unsupported file type: {file_ext}')

    def _extract_text(self, file_path: str, file_ext: str) -> str:
        """Extract text content from different file types"""
        if file_ext == '.pdf':
            return extract_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            doc = Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError(f'Unsupported file type: {file_ext}')

    async def process_resume(self, file_path: str, candidate_id: int = None) -> Dict[str, Any]:
        """Process resume file and save to database"""
        try:
            if not file_path or not os.path.exists(file_path):
                logging.error(f"Invalid or missing file path: {file_path}")
                return {"error": "Invalid or missing file path"}

            parsed_result = self.parse_file(file_path)
            if not parsed_result:
                logging.error(f"Failed to parse resume file: {file_path}")
                return {"error": "Failed to parse resume file"}

            logging.info(f"Resume parsed: {parsed_result}")

            return {"success": True, "message": "Resume processed successfully", "candidate_id": candidate_id, "data": parsed_result['parsed_data']}

        except Exception as e:
            logging.error(f"Error processing resume: {str(e)}")
            return {"error": f"Failed to process resume: {str(e)}"}